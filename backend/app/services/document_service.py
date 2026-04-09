import logging
import os
import uuid

from fastapi import UploadFile
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models.document import Document
from app.services import gigachat_client, vector_store

logger = logging.getLogger(__name__)

# Documents shorter than this are injected as full text in the prompt.
# Longer documents get chunked + embedded for similarity search.
CONTEXT_CHAR_LIMIT = 12_000

_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,
    chunk_overlap=200,
    separators=["\n\n", "\n", ". ", " ", ""],
)


def chunk_text(text: str) -> list[str]:
    """Split text into overlapping chunks for embedding."""
    return _splitter.split_text(text)


def is_large_document(text: str) -> bool:
    return len(text) > CONTEXT_CHAR_LIMIT


def index_document(document_id: str, text: str) -> int:
    """Chunk, embed, and store a document in the vector store.

    Returns the number of chunks created.
    """
    chunks = chunk_text(text)
    if not chunks:
        return 0
    # Embed in batches of 50 (GigaChat API limit)
    all_embeddings: list[list[float]] = []
    for i in range(0, len(chunks), 50):
        batch = chunks[i : i + 50]
        all_embeddings.extend(gigachat_client.get_embeddings(batch))
    vector_store.store_chunks(document_id, chunks, all_embeddings)
    logger.info("Indexed document %s: %d chunks", document_id, len(chunks))
    return len(chunks)


async def parse_text(file_path: str, content_type: str) -> str:
    """Extract plain text from a document file."""
    if content_type == "txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    elif content_type == "pdf":
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(pages)
    elif content_type == "docx":
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    else:
        raise ValueError(f"Unsupported content type: {content_type}")


def detect_content_type(filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    mapping = {"pdf": "pdf", "docx": "docx", "txt": "txt", "text": "txt"}
    if ext not in mapping:
        raise ValueError(f"Unsupported file extension: .{ext}")
    return mapping[ext]


async def upload_document(db: AsyncSession, file: UploadFile) -> Document:
    """Save uploaded file, extract text, store in DB."""
    filename = file.filename or "unnamed"
    content_type = detect_content_type(filename)

    # Save file to disk
    file_id = str(uuid.uuid4())
    ext = filename.rsplit(".", 1)[-1] if "." in filename else "bin"
    stored_name = f"{file_id}.{ext}"
    file_path = os.path.join(settings.upload_dir, stored_name)

    os.makedirs(settings.upload_dir, exist_ok=True)
    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)

    # Extract text
    extracted_text = await parse_text(file_path, content_type)

    # Save to DB
    document = Document(
        id=file_id,
        filename=filename,
        content_type=content_type,
        file_path=file_path,
        extracted_text=extracted_text,
        size_bytes=len(content),
    )
    db.add(document)
    await db.commit()
    await db.refresh(document)

    # Index large documents for similarity search
    if is_large_document(extracted_text):
        try:
            index_document(document.id, extracted_text)
        except Exception:
            logger.exception("Failed to index document %s", document.id)

    return document


async def upload_document_from_buffer(
    db: AsyncSession, filename: str, content: bytes
) -> Document:
    """Save document from pre-read buffer, extract text, store in DB."""
    content_type = detect_content_type(filename)

    file_id = str(uuid.uuid4())
    ext = filename.rsplit(".", 1)[-1] if "." in filename else "bin"
    stored_name = f"{file_id}.{ext}"
    file_path = os.path.join(settings.upload_dir, stored_name)

    os.makedirs(settings.upload_dir, exist_ok=True)
    with open(file_path, "wb") as f:
        f.write(content)

    extracted_text = await parse_text(file_path, content_type)

    document = Document(
        id=file_id,
        filename=filename,
        content_type=content_type,
        file_path=file_path,
        extracted_text=extracted_text,
        size_bytes=len(content),
    )
    db.add(document)
    await db.commit()
    await db.refresh(document)

    if is_large_document(extracted_text):
        try:
            index_document(document.id, extracted_text)
        except Exception:
            logger.exception("Failed to index document %s", document.id)

    return document


async def get_document(db: AsyncSession, document_id: str) -> Document | None:
    from sqlalchemy import select
    result = await db.execute(
        select(Document).where(Document.id == document_id)
    )
    return result.scalar_one_or_none()


async def delete_document(db: AsyncSession, document_id: str) -> bool:
    doc = await get_document(db, document_id)
    if not doc:
        return False
    if os.path.exists(doc.file_path):
        os.remove(doc.file_path)
    vector_store.delete_document_chunks(document_id)
    await db.delete(doc)
    await db.commit()
    return True
