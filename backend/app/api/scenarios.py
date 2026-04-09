import io
import json
from typing import Annotated

from fastapi import APIRouter, Depends, Form, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sse_starlette.sse import EventSourceResponse

from app.database import get_db
from app.models.scenario import ScenarioRun
from app.schemas.scenario import (
    ScenarioCreate,
    ScenarioDetailResponse,
    ScenarioResponse,
    ScenarioRunDetailResponse,
    ScenarioRunResponse,
    ScenarioUpdate,
)
from app.services import document_service, scenario_service
from app.services.scenario_engine import execute_scenario

router = APIRouter(prefix="/api/scenarios", tags=["scenarios"])


@router.get("", response_model=list[ScenarioResponse])
async def list_scenarios(db: AsyncSession = Depends(get_db)):
    return await scenario_service.get_scenarios(db)


@router.post("", response_model=ScenarioResponse, status_code=201)
async def create_scenario(body: ScenarioCreate, db: AsyncSession = Depends(get_db)):
    return await scenario_service.create_scenario(
        db, name=body.name, description=body.description, graph_data=body.graph_data
    )


@router.get("/runs/{run_id}", response_model=ScenarioRunDetailResponse)
async def get_scenario_run(run_id: str, db: AsyncSession = Depends(get_db)):
    run = await scenario_service.get_scenario_run(db, run_id)
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")
    return run


@router.get("/runs/{run_id}/export")
async def export_scenario_run(run_id: str, db: AsyncSession = Depends(get_db)):
    """Export scenario run result as DOCX."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor

    run = await scenario_service.get_scenario_run(db, run_id)
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")

    scenario = await scenario_service.get_scenario(db, run.scenario_id)

    doc = DocxDocument()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading(f'Результат: {scenario.name if scenario else "Сценарий"}', level=1)

    # Status
    doc.add_paragraph(f'Статус: {run.status}')
    if run.started_at:
        doc.add_paragraph(f'Запущен: {run.started_at}')
    if run.completed_at:
        doc.add_paragraph(f'Завершён: {run.completed_at}')

    doc.add_heading('Шаги выполнения', level=2)
    for step in run.steps:
        p = doc.add_heading(f'{step.node_id} ({step.node_type})', level=3)
        doc.add_paragraph(f'Статус: {step.status}')
        if step.output_data:
            output = step.output_data
            if isinstance(output, dict) and 'result' in output:
                doc.add_paragraph(output['result'])
            elif isinstance(output, dict) and 'error' in output:
                doc.add_paragraph(f'Ошибка: {output["error"]}')

    if run.result:
        doc.add_heading('Итоговый результат', level=2)
        output_text = run.result.get('output', '') if isinstance(run.result, dict) else str(run.result)
        doc.add_paragraph(output_text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"result_{run_id[:8]}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/{scenario_id}/duplicate", response_model=ScenarioResponse)
async def duplicate_scenario(scenario_id: str, db: AsyncSession = Depends(get_db)):
    copy = await scenario_service.duplicate_scenario(db, scenario_id)
    if not copy:
        raise HTTPException(status_code=404, detail="Scenario not found")
    return copy


@router.get("/{scenario_id}", response_model=ScenarioDetailResponse)
async def get_scenario(scenario_id: str, db: AsyncSession = Depends(get_db)):
    scenario = await scenario_service.get_scenario(db, scenario_id)
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    return scenario


@router.put("/{scenario_id}", response_model=ScenarioResponse)
async def update_scenario(
    scenario_id: str, body: ScenarioUpdate, db: AsyncSession = Depends(get_db)
):
    scenario = await scenario_service.update_scenario(
        db, scenario_id,
        name=body.name, description=body.description, graph_data=body.graph_data,
    )
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    return scenario


@router.delete("/{scenario_id}", status_code=204)
async def delete_scenario(scenario_id: str, db: AsyncSession = Depends(get_db)):
    deleted = await scenario_service.delete_scenario(db, scenario_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="Scenario not found")


@router.get("/{scenario_id}/runs", response_model=list[ScenarioRunResponse])
async def list_scenario_runs(scenario_id: str, db: AsyncSession = Depends(get_db)):
    return await scenario_service.get_scenario_runs(db, scenario_id)


@router.post("/{scenario_id}/run")
async def run_scenario(
    scenario_id: str,
    files: list[UploadFile] = [],
    db: AsyncSession = Depends(get_db),
):
    """Run a scenario with uploaded documents. Streams progress via SSE."""
    scenario = await scenario_service.get_scenario(db, scenario_id)
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")

    # Read file contents into memory before entering the generator
    # (UploadFile objects can't be read after the request handler returns)
    file_buffers: list[tuple[str, bytes]] = []
    for file in files:
        content = await file.read()
        file_buffers.append((file.filename or "unnamed", content))

    async def event_generator():
        try:
            # Phase 1: Upload and parse documents inside SSE stream
            documents_text_parts = []
            document_ids = []
            total_files = len(file_buffers)

            for i, (filename, content) in enumerate(file_buffers):
                yield {"data": json.dumps({
                    "type": "upload_progress",
                    "file": filename,
                    "current": i + 1,
                    "total": total_files,
                }, ensure_ascii=False)}

                try:
                    doc = await document_service.upload_document_from_buffer(
                        db, filename, content
                    )
                    document_ids.append(doc.id)
                    documents_text_parts.append(doc.extracted_text or "")
                except Exception as e:
                    yield {"data": json.dumps({
                        "type": "upload_error",
                        "file": filename,
                        "error": str(e),
                    }, ensure_ascii=False)}

            documents_text = "\n\n---\n\n".join(documents_text_parts)

            yield {"data": json.dumps({
                "type": "upload_done",
                "count": len(document_ids),
            }, ensure_ascii=False)}

            # Phase 2: Create run and execute scenario
            run = ScenarioRun(
                scenario_id=scenario_id,
                status="pending",
                input_document_ids=document_ids,
            )
            db.add(run)
            await db.commit()
            await db.refresh(run)

            async for event in execute_scenario(
                db, run, scenario.graph_data, documents_text
            ):
                yield {"data": json.dumps(event, ensure_ascii=False)}
        except Exception as e:
            yield {"data": json.dumps({"type": "error", "message": str(e)})}

    return EventSourceResponse(event_generator())
